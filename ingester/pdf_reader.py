import pdfplumber
from pathlib import Path
from utils.logger import get_logger

log = get_logger("pdf_reader")


def extract_text_from_pdf(path: Path) -> str:
    """Extract text from a PDF file."""
    full_text = []
    with pdfplumber.open(path) as pdf:
        for i, page in enumerate(pdf.pages):
            text = page.extract_text()
            if text:
                full_text.append(f"--- Page {i + 1} ---\n{text.strip()}")
    if not full_text:
        raise ValueError(f"Could not extract text from {path.name}. Is it a scanned PDF?")
    return "\n\n".join(full_text)


def extract_text_from_docx(path: Path) -> str:
    """Extract text from a Word .docx file."""
    from docx import Document
    doc = Document(str(path))
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    # Also extract tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                paragraphs.append(row_text)
    if not paragraphs:
        raise ValueError(f"Could not extract text from {path.name}")
    return "\n\n".join(paragraphs)


def extract_text(file_path: str) -> str:
    """Extract text from a PDF or DOCX file."""
    path = Path(file_path)

    if not path.exists():
        raise FileNotFoundError(
            f"File not found: {file_path}\n"
            f"Make sure it's placed in inputs/pdfs/ and named correctly in config.yaml"
        )

    log.info(f"Reading → {path.name}")
    suffix = path.suffix.lower()

    if suffix == ".pdf":
        text = extract_text_from_pdf(path)
    elif suffix in (".docx", ".doc"):
        text = extract_text_from_docx(path)
    else:
        raise ValueError(f"Unsupported file type: {suffix}. Use .pdf or .docx")

    log.info(f"Extracted {len(text):,} chars from {path.name}")
    return text


def load_all_pdfs(pdf_config: dict) -> dict:
    """
    Load all input files defined in config for a product.
    Supports .pdf and .docx. Skips missing optional files with a warning.

    Returns dict of { pdf_type: raw_text }
    """
    raw_texts = {}
    for pdf_type, file_path in pdf_config.items():
        path = Path(file_path)
        if not path.exists():
            log.warning(f"File not found, skipping: {file_path}")
            raw_texts[pdf_type] = f"[{pdf_type} not provided]"
            continue
        raw_texts[pdf_type] = extract_text(file_path)

    log.info(f"Loaded: {list(raw_texts.keys())}")
    return raw_texts
